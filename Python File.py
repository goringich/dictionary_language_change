from docx import Document
import re

# Load the original document
doc = Document("C:/Users/user/Desktop/rus_in_eng/1.docx")

# Create a new Word document
new_doc = Document()

# Define a regular expression pattern to match text within brackets
bracket_pattern = r'\([^)]*\)'

# Process each paragraph in the original document
for old_paragraph in doc.paragraphs:
    # Check the total length of words in the paragraph
    paragraph_length = sum(len(word) for word in old_paragraph.text.split())

    # Create a new paragraph in the new document
    new_paragraph = new_doc.add_paragraph()

    # Add a tab at the beginning of the paragraph if its length is greater than 5
    if paragraph_length > 5:
        new_paragraph.add_run('\t')

    # Split the text into words
    words = old_paragraph.text.split()

    # Initialize variables for tracking the style
    run_style = None
    in_brackets = False

    for word in words:
        # Check if the word is within brackets
        if re.search(bracket_pattern, word):
            in_brackets = not in_brackets

        if in_brackets:
            # If the word is within brackets, create a new run with the same style
            new_run = new_paragraph.add_run(" " + word)
            if run_style:
                new_run.bold = run_style.bold
                new_run.italic = run_style.italic
                new_run.underline = run_style.underline
                new_run.font.name = run_style.font.name
                new_run.font.size = run_style.font.size
                new_run.font.color.rgb = run_style.font.color.rgb
        else:
            # Create a new run for the word with the original style
            new_run = new_paragraph.add_run(" " + word)
            run_style = new_run

# Save the new document
new_doc.save("2_with_tabs.docx")

print("New document created with modified text and style.")
